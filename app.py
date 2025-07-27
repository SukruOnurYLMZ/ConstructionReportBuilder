import streamlit as st
from openai import OpenAI
from dotenv import load_dotenv
import os
from datetime import datetime
from docx import Document

# .env dosyasını yükle
load_dotenv()

# API anahtarları ve isim bilgileri
api_key = os.getenv("TOGETHER_API_KEY")
your_name = os.getenv("YOUR_NAME")
supervisor_name = os.getenv("SUPERVISOR_NAME")

# Together API istemcisi
client = OpenAI(
    api_key=api_key,
    base_url="https://api.together.xyz/v1"
)

# Rapor üretme fonksiyonu
def generate_report(prompt):
    response = client.chat.completions.create(
        model="meta-llama/Llama-3-70b-chat-hf",
        messages=[
            {"role": "system", "content": "Sen bir teknik şantiye raporu asistanısın. Verilen notlardan günlük mühendislik raporu üret."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.3,
    )
    return response.choices[0].message.content

# Word dökümanı oluşturma fonksiyonu
def create_docx(report, filename="report.docx"):
    doc = Document()
    doc.add_heading("Daily Construction Report", 0)
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph(f"Prepared By: {your_name}")
    doc.add_paragraph(f"Süpervizör: {supervisor_name}")
    doc.add_paragraph("\n" + report)
    doc.save(filename)
    return filename

# Streamlit arayüzü
st.title("📋 Construction Daily Report Form")
note = st.text_area("Write Your Daily Report For Today (include task hours and issues if possible:")

if st.button("Raporu Oluştur"):
    with st.spinner("Rapor oluşturuluyor..."):
        report = generate_report(note)
        filename = create_docx(report)
        st.success("✅ Report Completed!")
        with open(filename, "rb") as f:
            st.download_button(
                label="📥 Download Report(.docx)",
                data=f,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )